# creating docx file with title, heading, and paragraphs
# this is muc easier than creating Excel docs
import docx

document = docx.Document()

document.add_paragraph('Hello Word!', 'Title')
document.add_paragraph('My name is Johnny', 'Heading 1')
document.add_paragraph('This is a Word document created by Python using the python-docx library')
document.add_paragraph('Python Programming is a challenge!', 'BodyText')

document.add_paragraph('This paragraph has some ')
last_paragraph = document.paragraphs[-1]
last_paragraph.add_run('bold, underlined, and italicized text ')
last_run = last_paragraph.runs[-1]
last_run.bold = True
last_run.underline = True
last_run.italic = True
last_paragraph.add_run('in the middle')
document.save('hello_word.docx')

